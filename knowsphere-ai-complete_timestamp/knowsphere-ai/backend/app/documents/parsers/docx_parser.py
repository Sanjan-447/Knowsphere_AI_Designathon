from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.documents.parsers.base import BaseParser, ParsedDocument, ParsedSection, ParserError

# Word's built-in heading style names, used to detect section boundaries.
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}


class DocxParser(BaseParser):
    supported_extensions = ("docx",)

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            docx_doc = DocxDocument(file_path)
        except (PackageNotFoundError, KeyError, OSError) as exc:
            raise ParserError(f"Could not open DOCX '{file_path}': {exc}") from exc

        sections: list[ParsedSection] = []
        current_label: str | None = None
        current_lines: list[str] = []

        def flush():
            if current_lines:
                sections.append(ParsedSection(label=current_label, content="\n".join(current_lines).strip()))

        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style is not None and para.style.name in _HEADING_STYLES:
                flush()
                current_label = text
                current_lines = []
            else:
                current_lines.append(text)
        flush()

        # Tables aren't walked by .paragraphs — append them as their own section.
        for t_idx, table in enumerate(docx_doc.tables, start=1):
            rows_text = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            if rows_text:
                sections.append(ParsedSection(label=f"table {t_idx}", content="\n".join(rows_text)))

        core = docx_doc.core_properties
        metadata = {}
        if core.author:
            metadata["author"] = core.author
        if core.modified:
            metadata["source_last_modified"] = core.modified.isoformat()
        if core.title:
            metadata["docx_title"] = core.title

        doc = ParsedDocument(sections=sections, metadata=metadata)
        self.validate_not_empty(doc, file_path)
        return doc
